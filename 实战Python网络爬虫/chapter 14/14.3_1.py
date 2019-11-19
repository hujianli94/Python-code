# 数据写入
from docx import Document
from docx.shared import Inches
# 创建对象
document = Document()
# 添加标题，其中'0'代表标题类型，共四种类型，具体可在Word的'开始'-'样式'中查看
document.add_heading('Python 爬虫', 0)
# 添加正文内容并设置部分内容格式
p = document.add_paragraph('Python 爬虫开发-')
# 设置内容加粗
p.runs[0].bold = True
# 添加内容并加粗
p.add_run('数据存储-').bold = True
# 添加内容
p.add_run('Word-')
# 添加内容并设置字体斜体
p.add_run('存储实例。').italic = True
# 添加正文，设置'样式'-'明显引用'
document.add_paragraph('样式-明显引用', style='IntenseQuote')
# 添加正文，设置'项目符号'
document.add_paragraph(
    '项目符号1', style='ListBullet'
)
document.add_paragraph(
    '项目符号2', style='ListNumber'
)
# 添加图片
document.add_picture('test.png', width=Inches(1.25))
# 添加表格
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in range(2):
    row_cells = table.add_row().cells
    row_cells[0].text = 'a'
    row_cells[1].text = 'b'
    row_cells[2].text = 'c'
# 保存文件
document.add_page_break()
document.save('test.docx')
